"""
CODE AGENT - Main Streamlit Application
Turns natural language prompts + Figma links + tech stack choices into complete full-stack projects
"""

import streamlit as st
import sys
from pathlib import Path
from typing import Optional, Dict, Any
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent))

from workflow.orchestrator import ProjectOrchestrator
from utils.logger import StreamlitLogger
from utils.file_browser import build_file_tree, render_file_tree, preview_file, get_code_language, render_tree_visual, can_render_preview, render_file_preview
from utils.file_parser import FileParser
from utils.llm_factory import LLMFactory
# Removed ProjectRunner and ProjectPreview imports (no longer used)

# Page config
st.set_page_config(
    page_title="CODE AGENT - AI Full-Stack Platform",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for better UI
st.markdown("""
    <style>
    .main-header {
        font-size: 3.5rem;
        font-weight: bold;
        text-align: center;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 1rem;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
    }
    .sub-header {
        text-align: center;
        color: #666;
        margin-bottom: 2rem;
    }
    .stButton>button {
        width: 100%;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        font-size: 1.2rem;
        font-weight: bold;
        padding: 0.75rem 2rem;
        border-radius: 0.5rem;
        border: none;
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(102, 126, 234, 0.4);
    }
    </style>
""", unsafe_allow_html=True)

def main():
    # Project name at the top center
    st.markdown('<div style="text-align: center; font-size: 4rem; font-weight: bold; background: linear-gradient(90deg, #667eea 0%, #764ba2 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; margin-bottom: 0.5rem;">🚀 CODE AGENT</div>', unsafe_allow_html=True)
    
    st.markdown('<p class="sub-header" style="margin-top: 0;">AI-Powered Full-Stack Development Platform - Transform GitHub repositories into production-ready applications</p>', unsafe_allow_html=True)
    
    # Initialize session state
    if 'generation_in_progress' not in st.session_state:
        st.session_state.generation_in_progress = False
    if 'logs' not in st.session_state:
        st.session_state.logs = []
    if 'zip_file' not in st.session_state:
        st.session_state.zip_file = None
    if 'project_directory' not in st.session_state:
        st.session_state.project_directory = None
    if 'selected_file' not in st.session_state:
        st.session_state.selected_file = None
    if 'file_tree' not in st.session_state:
        st.session_state.file_tree = None
    if 'project_name' not in st.session_state:
        st.session_state.project_name = None
    if 'github_url' not in st.session_state:
        st.session_state.github_url = None
    if 'github_result' not in st.session_state:
        st.session_state.github_result = None
    # Removed preview and run project session state variables
    if 'last_project_config' not in st.session_state:
        st.session_state.last_project_config = None
    # Removed preview-related session state variables
    
    # Initialize form values first
    form_values = {
        "project_name": "",
        "backend_stack": "FastAPI + SQLAlchemy",
        "github_repo_url": ""
    }
    
    st.markdown("### 🐙 GitHub Repository")
    github_repo_url = st.text_input(
        "GitHub Repository URL *",
        value=form_values.get("github_repo_url", ""),
        placeholder="https://github.com/username/frontend-repo",
        help="GitHub repository URL containing the frontend code"
    )
    form_values["github_repo_url"] = github_repo_url
    
    # File upload section (MANDATORY)
    st.markdown("### 📁 Upload Files (Required)")
    st.error("⚠️ Impact Analysis report is MANDATORY to extract correct API endpoints for backend generation.")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**PRD File (Optional)**")
        prd_file = st.file_uploader(
            "Upload PRD (PDF, DOCX, TXT)",
            type=['pdf', 'docx', 'txt'],
            help="Optional: Upload your Product Requirements Document",
            key="prd_file"
        )
    
    with col2:
        st.markdown("**Impact Analysis Report (Required) ***")
        impact_file = st.file_uploader(
            "Upload Impact Analysis (PDF, DOCX, TXT) *",
            type=['pdf', 'docx', 'txt'],
            help="REQUIRED: Upload your Impact Analysis report containing API endpoints",
            key="impact_file"
        )
    
    # Parse uploaded files and display content below each file
    if prd_file is not None:
        with st.spinner(f"🔍 Analyzing PRD file: {prd_file.name}"):
            try:
                # Parse PRD file content
                if prd_file.name.endswith('.pdf'):
                    import PyPDF2
                    import io
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(prd_file.read()))
                    prd_content = ""
                    for page in pdf_reader.pages:
                        prd_content += page.extract_text() + "\n"
                elif prd_file.name.endswith('.docx'):
                    import docx
                    import io
                    doc = docx.Document(io.BytesIO(prd_file.read()))
                    prd_content = ""
                    for paragraph in doc.paragraphs:
                        prd_content += paragraph.text + "\n"
                else:
                    prd_content = prd_file.read().decode('utf-8')
                
                # Display content below PRD file
                with col1:
                    with st.expander(f"📄 PRD Content: {prd_file.name}", expanded=False):
                        st.text_area("PRD Content", prd_content[:2000] + "..." if len(prd_content) > 2000 else prd_content, height=200, disabled=True)
                        
            except Exception as e:
                with col1:
                    st.error(f"❌ Error reading PRD file: {str(e)}")
    
    if impact_file is not None:
        with st.spinner(f"🔍 Analyzing Impact Analysis: {impact_file.name}"):
            try:
                # Parse Impact Analysis file content
                if impact_file.name.endswith('.pdf'):
                    import PyPDF2
                    import io
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(impact_file.read()))
                    impact_content = ""
                    for page in pdf_reader.pages:
                        impact_content += page.extract_text() + "\n"
                elif impact_file.name.endswith('.docx'):
                    import docx
                    import io
                    doc = docx.Document(io.BytesIO(impact_file.read()))
                    impact_content = ""
                    for paragraph in doc.paragraphs:
                        impact_content += paragraph.text + "\n"
                else:
                    impact_content = impact_file.read().decode('utf-8')
                
                # Display content below Impact Analysis file
                with col2:
                    with st.expander(f"📄 Impact Analysis Content: {impact_file.name}", expanded=False):
                        st.text_area("Impact Analysis Content", impact_content[:2000] + "..." if len(impact_content) > 2000 else impact_content, height=200, disabled=True)
                        
            except Exception as e:
                with col2:
                    st.error(f"❌ Error reading Impact Analysis file: {str(e)}")
    
    st.markdown("Fill in the details manually or edit the extracted information:")
    
    # Main form
    with st.form("project_generator_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            project_name = st.text_input(
                "Project Name *",
                value=form_values["project_name"],
                placeholder="my-awesome-app",
                help="Name for your project (will be used for folder structure)"
            )
            
            # Project description removed - will be extracted from files
            
            # GitHub URL is now in the upload section above
            pass
        
        with col2:
            backend_options = ["FastAPI + SQLAlchemy", "Django", "Node.js/Express + Prisma", "Supabase", "Firebase"]
            backend_index = 0
            if form_values["backend_stack"] in backend_options:
                backend_index = backend_options.index(form_values["backend_stack"])
            
            backend_stack = st.selectbox(
                "Backend Stack *",
                options=backend_options,
                index=backend_index,
                help="Choose your backend framework"
            )
            
            # Remove Figma input - not needed anymore
            pass
            
            # GitHub publishing option
            publish_to_github = st.checkbox("🐙 Publish to GitHub", help="Automatically create a GitHub repository and push your project")
        
        submitted = st.form_submit_button(
            "🚀 Generate Full Project",
            use_container_width=True
        )
    
    # Handle form submission
    if submitted and not st.session_state.generation_in_progress:
        # Validation - GitHub URL, project name, and Impact Analysis are required
        if not github_repo_url:
            st.error("Please provide a GitHub repository URL (marked with *)")
            return
        if not project_name:
            st.error("Please provide a project name (marked with *)")
            return
        if not impact_file:
            st.error("Please upload Impact Analysis report (marked with *) - This is MANDATORY for correct API endpoint extraction")
            return
        
        # Use default API key
        api_key = None
        
        st.session_state.generation_in_progress = True
        st.session_state.logs = []
        st.session_state.zip_file = None
        st.session_state.project_directory = None
        st.session_state.selected_file = None
        st.session_state.file_tree = None
        st.session_state.project_name = project_name
        # Removed project run state variables
        st.session_state.last_project_config = None
        
        # Validate GitHub URL
        if not github_repo_url.startswith('https://github.com/'):
            st.error("Please provide a valid GitHub repository URL")
            return
        
        # GitHub URL validated, proceed with generation
        st.success(f"✅ GitHub repository URL validated: {github_repo_url}")
        
        # Initialize logger
        logger = StreamlitLogger()
        
        # Show progress area
        st.markdown("---")
        st.markdown("### 📊 Generation Progress")
        log_container = st.empty()
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        try:
            # Initialize orchestrator with default API key
            selected_api_key = None
            orchestrator = ProjectOrchestrator(
                api_key=selected_api_key,  # Use custom key if provided, otherwise None (uses default)
                logger=logger
            )
            
            # Prepare project config
            project_config = {
                "project_name": project_name,
                "description": f"Backend API for {project_name}",  # Default description
                "frontend_stack": "React",  # Default frontend stack since we're cloning from GitHub
                "backend_stack": backend_stack,
                "github_repo_url": github_repo_url,
                "publish_to_github": publish_to_github,
                "prd_file_content": prd_file.getvalue() if prd_file else None,
                "prd_file_name": prd_file.name if prd_file else None,
                "impact_file_content": impact_file.getvalue() if impact_file else None,
                "impact_file_name": impact_file.name if impact_file else None
            }
            st.session_state.last_project_config = project_config
            
            # Run generation workflow
            with st.spinner("Initializing project generation..."):
                result = orchestrator.generate_project(project_config)
            
            # Update UI with logs
            update_logs_display(log_container, logger.get_logs())
            progress_bar.progress(1.0)
            status_text.success("✅ Project generated successfully!")
            
            # Get ZIP file path and project directory from result
            if result and result.get("project_path"):
                zip_path = result["project_path"]
                project_dir = result.get("project_directory")
                github_result = result.get("github_result")
                
                # Store project directory in session state
                if project_dir:
                    st.session_state.project_directory = project_dir
                
                # Store GitHub result and URL in session state
                if github_result:
                    st.session_state.github_result = github_result
                    if github_result.get("success"):
                        repo_url = github_result.get("repository_url")
                        if repo_url:
                            st.session_state.github_url = repo_url
                
                # Read ZIP file
                try:
                    with open(zip_path, "rb") as f:
                        zip_buffer = f.read()
                    st.session_state.zip_file = zip_buffer
                except Exception as e:
                    st.warning(f"ZIP file created at: {zip_path}")
                    st.error(f"Could not read ZIP file: {str(e)}")
            
        except Exception as e:
            error_message = str(e)
            logger.log(f"ERROR: {error_message}", level="error")
            update_logs_display(log_container, logger.get_logs())
            
            # Check for insufficient credits/payment required errors
            if "402" in error_message or "requires more credits" in error_message.lower() or "can only afford" in error_message.lower():
                st.error("💰 **Insufficient Credits Error**")
                st.warning(f"""
                **Your OpenRouter account doesn't have enough credits for this request.**
                
                **The Issue:**
                - You requested up to 16384 tokens, but can only afford 4000 tokens
                - This is a free tier limitation
                
                **Solutions:**
                1. ✅ **Upgrade Your OpenRouter Account** (Recommended)
                   - Visit: https://openrouter.ai/settings/credits
                   - Add credits to your account
                   - This allows you to use premium models like GPT-4o
                
                2. ✅ **Use a Free Model** (Quick Fix)
                   - The app will automatically use a free model (Llama 3.1 70B) for free tier
                   - This model is free but may be slower
                   - Try generating again - it should work now
                
                3. ✅ **Reduce Project Complexity**
                   - Try a smaller project description
                   - The app now limits to 4000 tokens to stay within free tier
                
                **Note:** The app has been updated to use free-tier compatible models and token limits.
                """)
            # Check for invalid API key errors
            elif "401" in error_message or "invalid_api_key" in error_message.lower() or "Invalid API Key" in error_message:
                st.error("🔑 **Invalid API Key Error**")
                st.warning(f"""
                **Your OpenRouter API key is invalid or expired.**
                
                **Quick Fix (Easiest):**
                1. ✅ **Get a new API key** from https://openrouter.ai/keys
                   - Sign up or log in
                   - Go to **API Keys** section
                   - Click **"Create Key"**
                   - Copy the key (starts with `sk-or-v1-`)
                
                2. ✅ **Use it in the app:**
                   - Check the **"Use custom OpenRouter API key"** checkbox above
                   - Paste your new key in the field
                   - Click **"🚀 Generate Full Project"** again
                
                **Alternative: Update .env file:**
                - Open `grok/.env` file
                - Replace the `OPENROUTER_API_KEY` value with your new key
                - Restart the Streamlit app
                
                **Note:** The default API key may have expired. Using your own key is recommended!
                """)
            # Check for rate limit errors
            elif "rate_limit" in error_message.lower() or "429" in error_message or "Rate limit" in error_message or "tokens per day" in error_message.lower():
                st.error("⚠️ **API Rate Limit Reached**")
                
                # Try to extract information from error message
                import re
                reset_match = re.search(r'Please try again in ([\d\.]+[smh])', error_message)
                limit_match = re.search(r'Limit (\d+)', error_message)
                
                st.warning(f"""
                **Your OpenRouter API key has reached its rate limit.**
                
                {f"⏰ **Reset time:** {reset_match.group(1)}" if reset_match else ""}
                {f"📊 **Daily limit:** {limit_match.group(1)} tokens" if limit_match else ""}
                
                **Solutions:**
                1. **⏳ Wait** - The limit will reset automatically
                2. **🔑 Use a different API key** - Get a key from https://openrouter.ai/keys
                   - Check "Use custom OpenRouter API key" in the form above
                   - Paste your new key
                3. **💎 Upgrade** - Upgrade your OpenRouter plan for higher limits at https://openrouter.ai/settings
                """)
            else:
                st.error(f"❌ Error during generation: {error_message}")
            
        finally:
            st.session_state.generation_in_progress = False
    
    # Display project files browser and preview
    if st.session_state.project_directory:
        project_dir_path = Path(st.session_state.project_directory)
        
        if project_dir_path.exists():
            st.markdown("---")
            st.markdown("## 📂 Project Files")
            
            # GitHub Results Display (if available)
            if hasattr(st.session_state, 'github_result') and st.session_state.github_result:
                github_result = st.session_state.github_result
                if github_result.get("success"):
                    st.success("🎉 Project published to GitHub successfully!")
                    repo_url = github_result.get("repository_url")
                    if repo_url:
                        st.markdown(f"🔗 **Repository URL:** [{repo_url}]({repo_url})")
                        st.markdown(f"📁 **Files uploaded:** {github_result.get('uploaded_files', 0)}")
                elif github_result.get("skip_reason") == "no_token":
                    st.info("💡 **GitHub publishing skipped** - No GitHub token configured")
                else:
                    error_msg = github_result.get('error', 'Unknown error')
                    # Check for authentication errors
                    if "401" in str(error_msg) or "Unauthorized" in str(error_msg) or "invalid" in str(error_msg).lower():
                        st.error("🔑 **GitHub Authentication Failed**")
                        st.warning(f"""
                        **Your GitHub token is invalid, expired, or missing required permissions.**
                        
                        **Quick Fix:**
                        1. ✅ **Get a new GitHub Personal Access Token:**
                           - Go to: https://github.com/settings/tokens
                           - Click **"Generate new token"** → **"Generate new token (classic)"**
                           - Give it a name (e.g., "CODE AGENT")
                           - Select scopes: ✅ **`repo`** (Full control of private repositories)
                           - Click **"Generate token"**
                           - **Copy the token immediately** (starts with `ghp_`)
                        
                        2. ✅ **Add it to your `.env` file:**
                           - Open `grok/.env` file
                           - Add or update: `GITHUB_ACCESS_TOKEN=ghp_your_token_here`
                           - Save the file
                        
                        3. ✅ **Restart the Streamlit app** for changes to take effect
                        
                        **Note:** The token needs `repo` scope to create and push to repositories.
                        """)
                    else:
                        st.warning(f"⚠️ GitHub publishing failed: {error_msg}")
            
            # Main action buttons
            col1, col2 = st.columns(2)
            
            with col1:
                if st.session_state.zip_file:
                    st.download_button(
                        label="📦 Download ZIP",
                        data=st.session_state.zip_file,
                        file_name=f"{st.session_state.project_name or 'project'}.zip",
                        mime="application/zip",
                        use_container_width=True
                    )
            
            with col2:
                # Show GitHub link if available
                if hasattr(st.session_state, 'github_url') and st.session_state.github_url:
                    st.link_button(
                        "🐙 View on GitHub",
                        st.session_state.github_url,
                        use_container_width=True
                    )

            # Removed Preview Project and Run Project buttons and their functionality
            
            # Build file tree if not already built
            if st.session_state.file_tree is None:
                st.session_state.file_tree = build_file_tree(project_dir_path)
            
            # Removed Project Preview Section
            
            # Removed Project Preview Section
            
            # Two column layout: file browser and preview
            col1, col2 = st.columns([1, 2])
            
            with col1:
                st.markdown("### 📁 File Explorer")
                
                # File selector dropdown
                clicked_file = render_file_tree(st.session_state.file_tree, project_dir_path, st.session_state.selected_file)
                
                if clicked_file:
                    st.session_state.selected_file = clicked_file
                
                # Show tree structure
                st.markdown("---")
                st.markdown("**📂 Project Structure:**")
                render_tree_visual(st.session_state.file_tree, project_dir_path)
            
            with col2:
                st.markdown("### 👁️ File Preview")
                
                if st.session_state.selected_file:
                    selected_path = Path(st.session_state.selected_file)
                    
                    if selected_path.exists() and selected_path.is_file():
                        file_name = selected_path.name
                        file_size = selected_path.stat().st_size
                        relative_path = selected_path.relative_to(project_dir_path)
                        
                        # File metadata
                        col_path, col_size = st.columns([3, 1])
                        with col_path:
                            st.markdown(f"**📄 File:** `{relative_path}`")
                        with col_size:
                            st.markdown(f"**Size:** {file_size:,} bytes")
                        
                        # Check if file can be rendered as preview
                        if can_render_preview(selected_path):
                            # Create tabs for code view and live preview
                            tab1, tab2 = st.tabs(["📝 Code", "🌐 Live Preview"])
                            
                            with tab1:
                                # Show code with syntax highlighting
                                content = preview_file(selected_path)
                                language = get_code_language(selected_path)
                                
                                if content:
                                    st.code(content, language=language, line_numbers=True)
                                else:
                                    st.warning("Could not preview this file type")
                            
                            with tab2:
                                # Show live preview
                                try:
                                    preview_content, render_type = render_file_preview(selected_path)
                                    
                                    if render_type == 'html':
                                        st.markdown("**🌐 HTML Preview:**")
                                        st.components.v1.html(preview_content, height=600, scrolling=True)
                                    elif render_type == 'svg':
                                        st.markdown("**🖼️ SVG Preview:**")
                                        st.image(preview_content)
                                    elif render_type == 'markdown':
                                        st.markdown("**📝 Markdown Preview:**")
                                        st.markdown(preview_content)
                                    else:
                                        st.warning("Live preview not available for this file type")
                                        
                                except Exception as e:
                                    st.error(f"Error rendering preview: {str(e)}")
                        else:
                            # Show only code view for non-previewable files
                            content = preview_file(selected_path)
                            language = get_code_language(selected_path)
                            
                            if content:
                                st.code(content, language=language, line_numbers=True)
                            else:
                                st.warning("Could not preview this file type")
                    else:
                        st.info("Please select a file from the explorer to preview")
                else:
                    st.info("👈 Click on a file in the explorer to preview it here")
    
    # Display logs if available
    if st.session_state.logs:
        with st.expander("📋 View Generation Logs", expanded=False):
            for log in st.session_state.logs[-50:]:  # Show last 50 logs
                st.text(log)

# Removed Figma functions - no longer needed

def update_logs_display(container, logs):
    """Update the log display container"""
    if logs:
        log_text = "\n".join([f"[{log.get('timestamp', '')}] {log.get('message', '')}" for log in logs[-30:]])
        container.text_area("Live Logs", log_text, height=400, disabled=True, key="live_logs")

# Removed create_zip_file - packager agent handles ZIP creation

if __name__ == "__main__":
    main()

