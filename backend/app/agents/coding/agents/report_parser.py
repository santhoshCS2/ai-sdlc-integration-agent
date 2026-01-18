"""
ReportParserAgent - Reads and understands PRD and Impact Analysis reports
"""

from typing import Dict, Any, List
from langchain_core.prompts import ChatPromptTemplate
import json
from app.agents.coding.utils.logger import StreamlitLogger

class ReportParserAgent:
    """Agent that reads and understands PRD/Impact Analysis reports"""
    
    def __init__(self, llm, logger: StreamlitLogger):
        self.llm = llm
        self.logger = logger
    
    def read_report(self, project_config: Dict[str, Any]) -> Dict[str, Any]:
        """Read and understand PRD/Impact Analysis report content"""
        self.logger.log("📄 Reading PRD/Impact Analysis reports...")
        
        prd_content = project_config.get("prd_file_content")
        prd_name = project_config.get("prd_file_name", "")
        impact_content = project_config.get("impact_file_content")
        impact_name = project_config.get("impact_file_name", "")
        
        if not prd_content and not impact_content:
            self.logger.log("⚠️ No report files provided")
            return {"content": "", "summary": "No reports provided"}
        
        try:
            combined_text = ""
            
            # Parse PRD file if provided
            if prd_content:
                self.logger.log(f"📄 Processing PRD file: {prd_name}")
                if prd_name.endswith('.pdf'):
                    prd_text = self._parse_pdf(prd_content)
                elif prd_name.endswith('.docx'):
                    prd_text = self._parse_docx(prd_content)
                else:
                    prd_text = prd_content.decode('utf-8') if isinstance(prd_content, bytes) else str(prd_content)
                
                combined_text += f"PRD DOCUMENT:\n{prd_text}\n\n"
            
            # Parse Impact Analysis file if provided
            if impact_content:
                self.logger.log(f"📄 Processing Impact Analysis: {impact_name}")
                if impact_name.endswith('.pdf'):
                    impact_text = self._parse_pdf(impact_content)
                elif impact_name.endswith('.docx'):
                    impact_text = self._parse_docx(impact_content)
                else:
                    impact_text = impact_content.decode('utf-8') if isinstance(impact_content, bytes) else str(impact_content)
                
                combined_text += f"IMPACT ANALYSIS REPORT:\n{impact_text}\n\n"
            
            if not combined_text:
                return {"content": "", "summary": "No content extracted"}
            
            # Understand the full report content
            report_analysis = self._analyze_report_content(combined_text)
            
            self.logger.log(f"✅ Successfully read and analyzed report content ({len(combined_text)} characters)")
            return {
                "content": combined_text,
                "analysis": report_analysis,
                "prd_included": bool(prd_content),
                "impact_included": bool(impact_content)
            }
            
        except Exception as e:
            self.logger.log(f"⚠️ Error reading reports: {str(e)}", level="warning")
            return {"content": "", "summary": f"Error: {str(e)}"}
    
    def _parse_pdf(self, file_content: bytes) -> str:
        """Parse PDF file content with enhanced extraction"""
        try:
            import PyPDF2
            import io
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            
            self.logger.log(f"📝 Extracting text from {len(pdf_reader.pages)} PDF pages...")
            
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- PAGE {i+1} ---\n{page_text}\n"
                else:
                    self.logger.log(f"⚠️ Page {i+1} appears to be empty or image-based")
            
            self.logger.log(f"✅ Extracted {len(text)} characters from PDF")
            return text
            
        except Exception as e:
            self.logger.log(f"⚠️ Error parsing PDF: {str(e)}", level="warning")
            return ""
    
    def _parse_docx(self, file_content: bytes) -> str:
        """Parse DOCX file content with enhanced extraction"""
        try:
            import docx
            import io
            
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            
            self.logger.log(f"📝 Extracting text from DOCX with {len(doc.paragraphs)} paragraphs...")
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                text += "\n--- TABLE ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "--- END TABLE ---\n"
            
            self.logger.log(f"✅ Extracted {len(text)} characters from DOCX")
            return text
            
        except Exception as e:
            self.logger.log(f"⚠️ Error parsing DOCX: {str(e)}", level="warning")
            return ""
    
    def _analyze_report_content(self, text_content: str) -> Dict[str, Any]:
        """Analyze and understand the full report content"""
        self.logger.log(f"📊 Analyzing {len(text_content)} characters of report content...")
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert business analyst. Read and understand the complete document content including:

- Project specifications and requirements
- Backend architecture and structure
- Technical implementation details
- Business logic and workflows
- Data models and relationships
- Integration requirements
- Security considerations
- Performance requirements

Provide a comprehensive analysis in JSON format:
{
  "project_overview": "Brief summary of the project",
  "key_requirements": ["requirement1", "requirement2"],
  "backend_structure": "Description of backend architecture",
  "technical_details": "Key technical implementation points",
  "data_models": "Information about data structures",
  "integrations": "External systems or APIs mentioned",
  "security_notes": "Security requirements or considerations",
  "performance_notes": "Performance requirements or considerations"
}

Read the ENTIRE document thoroughly and extract ALL relevant information."""),
            ("human", "Analyze this complete document and provide comprehensive understanding:\n\n{text_content}")
        ])
        
        try:
            self.logger.log(f"🔍 Content sample (first 500 chars): {text_content[:500]}...")
            
            messages = prompt.format_messages(text_content=text_content)
            response = self.llm.invoke(messages)
            content = response.content.strip()
            
            # Parse JSON response
            analysis = self._parse_json_response(content)
            
            if analysis:
                self.logger.log("✅ Successfully analyzed report content")
                return analysis
            else:
                return {"summary": "Content read but analysis failed"}
                
        except Exception as e:
            self.logger.log(f"⚠️ Error analyzing content: {str(e)}", level="warning")
            return {"summary": f"Analysis error: {str(e)}"}
    
    def _parse_json_response(self, content: str) -> Dict[str, Any]:
        """Parse JSON response from LLM with high robustness"""
        try:
            import re
            # Clean up the response
            content = content.strip()
            
            # Method 1: Extract from markdown code blocks
            json_match = re.search(r'```(?:json)?\s*(\{[\s\S]*?\})\s*```', content, re.DOTALL)
            if json_match:
                content = json_match.group(1)
            else:
                # Method 2: Find JSON object directly
                json_match = re.search(r'\{[\s\S]*\}', content)
                if json_match:
                    content = json_match.group(0)
            
            return json.loads(content.strip())
        except json.JSONDecodeError as e:
            self.logger.log(f"⚠️ JSON parse error: {str(e)}", level="warning")
            return None
        except Exception as e:
            self.logger.log(f"⚠️ Response parse error: {str(e)}", level="warning")
            return None